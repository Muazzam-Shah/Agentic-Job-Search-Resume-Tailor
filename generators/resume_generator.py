"""
Resume Generator - Creates tailored, ATS-optimized resumes from parsed data

This module generates professional DOCX resumes by:
1. Selecting most relevant content based on job keywords
2. Optimizing bullet points with GPT-4o-mini
3. Applying ATS-friendly formatting
4. Generating properly named output files
"""

import os
import re
from typing import List, Dict, Optional, Tuple
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pydantic import BaseModel, Field
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import ChatPromptTemplate

from parsers.resume_parser import ParsedResume, Experience, Education
from tools.keyword_extractor import KeywordExtractor, ExtractedKeywords
from tools.semantic_matcher import SemanticMatcher
from utils.logger import logger



class OptimizedBulletPoint(BaseModel):
    """Optimized bullet point with job-specific keywords"""
    original: str = Field(description="Original bullet point text")
    optimized: str = Field(description="Optimized bullet point with keywords")
    keywords_added: List[str] = Field(description="Keywords incorporated")
    impact_score: float = Field(description="Estimated impact score (0-1)")


class OptimizedExperience(BaseModel):
    """Optimized experience section"""
    bullet_points: List[OptimizedBulletPoint] = Field(description="Optimized bullets")
    relevance_score: float = Field(description="Relevance to job (0-1)")
    selected_count: int = Field(description="Number of bullets to include")


class ContentSelection(BaseModel):
    """Selected content for tailored resume"""
    experiences: List[Tuple[Experience, List[str]]] = Field(
        description="Selected experiences with optimized bullets"
    )
    skills: List[str] = Field(description="Prioritized skills list")
    summary: str = Field(description="Tailored professional summary")
    selected_keywords: List[str] = Field(description="Keywords to emphasize")


class ResumeGenerator:
    """
    Generates tailored, ATS-optimized resumes
    
    Features:
    - Content selection based on job match
    - Bullet point optimization with GPT-4o-mini
    - ATS-friendly DOCX formatting
    - Keyword density optimization
    - Professional templates
    """
    
    def __init__(
        self,
        model: str = "gpt-4o-mini",
        temperature: float = 0.7,
        output_dir: str = "output/resumes"
    ):
        """
        Initialize resume generator
        
        Args:
            model: LLM model for bullet point optimization
            temperature: Temperature for creative rewriting (0.7 for balance)
            output_dir: Directory for generated resumes
        """
        self.model = model
        self.temperature = temperature
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize LLM for bullet point optimization
        self.llm = ChatOpenAI(
            model=model,
            temperature=temperature,
            model_kwargs={"response_format": {"type": "json_object"}}
        )
        
        # Initialize helper tools
        self.keyword_extractor = KeywordExtractor(model=model)
        self.semantic_matcher = SemanticMatcher()
        
        logger.info(f"ResumeGenerator initialized with {model}")
    
    def generate_tailored_resume(
        self,
        parsed_resume: ParsedResume,
        job_description: str,
        company_name: str,
        job_title: str,
        match_analysis: Optional[Dict] = None
    ) -> str:
        """
        Generate a tailored resume for specific job
        
        Args:
            parsed_resume: Parsed master resume data
            job_description: Target job description
            company_name: Company name for file naming
            job_title: Job title for file naming
            match_analysis: Pre-computed match analysis (optional)
        
        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Generating tailored resume for {company_name} - {job_title}")
        
        # Extract job keywords
        job_keywords = self.keyword_extractor.extract_keywords(job_description)
        
        # Get match analysis if not provided
        if match_analysis is None:
            match_analysis = self.semantic_matcher.analyze_match(
                parsed_resume.raw_text,
                parsed_resume,
                job_description
            )
        
        # Select and optimize content
        selected_content = self._select_content(
            parsed_resume,
            job_keywords,
            match_analysis
        )
        
        # Create tailored summary
        tailored_summary = self._create_tailored_summary(
            parsed_resume,
            job_keywords,
            job_description
        )
        
        # Optimize experiences
        optimized_experiences = self._optimize_experiences(
            selected_content['experiences'],
            job_keywords,
            job_description
        )
        
        # Generate DOCX
        output_path = self._generate_docx(
            parsed_resume=parsed_resume,
            tailored_summary=tailored_summary,
            optimized_experiences=optimized_experiences,
            selected_skills=selected_content['skills'],
            company_name=company_name,
            job_title=job_title
        )
        
        logger.info(f"Resume generated: {output_path}")
        return output_path
    
    def _select_content(
        self,
        parsed_resume: ParsedResume,
        job_keywords: ExtractedKeywords,
        match_analysis: Dict
    ) -> Dict:
        """
        Select most relevant content based on job requirements
        
        Args:
            parsed_resume: Parsed resume data
            job_keywords: Extracted job keywords
            match_analysis: Match analysis results
        
        Returns:
            Dictionary with selected content
        """
        logger.info("Selecting relevant content for job")
        
        # Score each experience by relevance
        experience_scores = []
        for exp in parsed_resume.experience:
            exp_text = f"{exp.title} {exp.company} {' '.join(exp.description or [])}"
            
            # Calculate keyword overlap
            all_keywords = (
                job_keywords.required_skills +
                job_keywords.preferred_skills +
                job_keywords.tools_technologies
            )
            
            keyword_matches = sum(
                1 for kw in all_keywords
                if kw.lower() in exp_text.lower()
            )
            
            relevance_score = keyword_matches / max(len(all_keywords), 1)
            experience_scores.append((exp, relevance_score))
        
        # Sort by relevance
        experience_scores.sort(key=lambda x: x[1], reverse=True)
        
        # Select top experiences (at least top 3, all if score > 0.2)
        selected_experiences = []
        for exp, score in experience_scores:
            if len(selected_experiences) < 3 or score > 0.2:
                selected_experiences.append((exp, exp.description or []))
        
        # Prioritize skills based on job keywords
        selected_skills = self._prioritize_skills(
            parsed_resume.skills,
            job_keywords
        )
        
        return {
            'experiences': selected_experiences,
            'skills': selected_skills,
            'experience_scores': experience_scores
        }
    
    def _prioritize_skills(
        self,
        resume_skills: List[str],
        job_keywords: ExtractedKeywords
    ) -> List[str]:
        """
        Prioritize skills based on job requirements
        
        Args:
            resume_skills: List of skills from resume
            job_keywords: Job keywords
        
        Returns:
            Prioritized skill list
        """
        if not resume_skills:
            return []
        
        # Create priority tiers
        required = set(kw.lower() for kw in job_keywords.required_skills)
        preferred = set(kw.lower() for kw in job_keywords.preferred_skills)
        tools = set(kw.lower() for kw in job_keywords.tools_technologies)
        
        # Categorize resume skills
        priority_1 = []  # Required skills
        priority_2 = []  # Tools and preferred
        priority_3 = []  # Other skills
        
        for skill in resume_skills:
            skill_lower = skill.lower()
            if skill_lower in required:
                priority_1.append(skill)
            elif skill_lower in preferred or skill_lower in tools:
                priority_2.append(skill)
            else:
                priority_3.append(skill)
        
        # Combine with priorities first
        return priority_1 + priority_2 + priority_3[:10]  # Limit total
    
    def _create_tailored_summary(
        self,
        parsed_resume: ParsedResume,
        job_keywords: ExtractedKeywords,
        job_description: str
    ) -> str:
        """
        Create job-specific professional summary using GPT-4o-mini
        
        Args:
            parsed_resume: Parsed resume
            job_keywords: Job keywords
            job_description: Full job description
        
        Returns:
            Tailored summary text
        """
        logger.info("Creating tailored professional summary")
        
        # Build prompt for summary generation
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert resume writer. Create a compelling professional summary 
            tailored to the specific job. The summary should:
            1. Highlight relevant skills and experience matching the job
            2. Incorporate key job requirements naturally
            3. Be concise (2-3 sentences, ~50 words)
            4. Use action-oriented language
            5. Emphasize quantifiable achievements when possible
            
            Return JSON: {{"summary": "tailored summary text"}}"""),
            ("user", """Original Summary: {original_summary}
            
Job Title/Description: {job_description}

Required Skills: {required_skills}
Preferred Skills: {preferred_skills}
Tools/Technologies: {tools}

Create a tailored professional summary that positions the candidate perfectly for this role.""")
        ])
        
        try:
            chain = prompt | self.llm | JsonOutputParser()
            result = chain.invoke({
                "original_summary": parsed_resume.summary or "Professional with diverse experience",
                "job_description": job_description[:500],  # Truncate for token limit
                "required_skills": ", ".join(job_keywords.required_skills[:10]),
                "preferred_skills": ", ".join(job_keywords.preferred_skills[:5]),
                "tools": ", ".join(job_keywords.tools_technologies[:5])
            })
            
            return result.get('summary', parsed_resume.summary or "")
            
        except Exception as e:
            logger.error(f"Failed to generate tailored summary: {e}")
            # Fallback to original summary
            return parsed_resume.summary or ""
    
    def _optimize_experiences(
        self,
        experiences: List[Tuple[Experience, List[str]]],
        job_keywords: ExtractedKeywords,
        job_description: str
    ) -> List[Tuple[Experience, List[str]]]:
        """
        Optimize experience bullet points with job keywords
        
        Args:
            experiences: List of (Experience, bullet_points) tuples
            job_keywords: Job keywords
            job_description: Job description
        
        Returns:
            List of (Experience, optimized_bullets) tuples
        """
        logger.info(f"Optimizing {len(experiences)} experience sections")
        
        optimized = []
        for exp, bullets in experiences:
            if not bullets:
                optimized.append((exp, []))
                continue
            
            # Optimize bullets
            optimized_bullets = self._optimize_bullet_points(
                bullets,
                job_keywords,
                exp.title
            )
            
            optimized.append((exp, optimized_bullets))
        
        return optimized
    
    def _optimize_bullet_points(
        self,
        bullets: List[str],
        job_keywords: ExtractedKeywords,
        role_title: str
    ) -> List[str]:
        """
        Optimize bullet points using GPT-4o-mini
        
        Args:
            bullets: Original bullet points
            job_keywords: Job keywords to incorporate
            role_title: Role title for context
        
        Returns:
            Optimized bullet points
        """
        if not bullets:
            return []
        
        # Build optimization prompt
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert resume writer. Optimize bullet points to:
            1. Incorporate relevant job keywords naturally
            2. Use strong action verbs
            3. Quantify achievements when possible
            4. Keep concise (1-2 lines each)
            5. Maintain truthfulness (enhance, don't fabricate)
            6. Use past tense for completed roles
            
            Return JSON: {{"optimized_bullets": ["bullet 1", "bullet 2", ...]}}"""),
            ("user", """Role: {role_title}

Original Bullets:
{bullets}

Job Keywords to Incorporate:
Required Skills: {required_skills}
Tools/Technologies: {tools}
Preferred Skills: {preferred_skills}

Optimize these bullet points to better match the job requirements while staying truthful.""")
        ])
        
        try:
            chain = prompt | self.llm | JsonOutputParser()
            result = chain.invoke({
                "role_title": role_title,
                "bullets": "\n".join(f"- {b}" for b in bullets),
                "required_skills": ", ".join(job_keywords.required_skills[:8]),
                "tools": ", ".join(job_keywords.tools_technologies[:6]),
                "preferred_skills": ", ".join(job_keywords.preferred_skills[:4])
            })
            
            optimized = result.get('optimized_bullets', bullets)
            logger.info(f"Optimized {len(bullets)} bullets for {role_title}")
            return optimized
            
        except Exception as e:
            logger.error(f"Failed to optimize bullets: {e}")
            return bullets  # Fallback to original
    
    def _generate_docx(
        self,
        parsed_resume: ParsedResume,
        tailored_summary: str,
        optimized_experiences: List[Tuple[Experience, List[str]]],
        selected_skills: List[str],
        company_name: str,
        job_title: str
    ) -> str:
        """
        Generate ATS-friendly DOCX resume
        
        Args:
            parsed_resume: Original parsed resume
            tailored_summary: Tailored summary
            optimized_experiences: Optimized experiences
            selected_skills: Prioritized skills
            company_name: Company name
            job_title: Job title
        
        Returns:
            Path to generated file
        """
        logger.info("Generating DOCX file")
        
        # Create document
        doc = Document()
        
        # Set document margins (1 inch all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add header with name and contact
        self._add_header(doc, parsed_resume.contact_info)
        
        # Add professional summary (renamed to OBJECTIVE for consistency)
        if tailored_summary:
            self._add_section(doc, "OBJECTIVE", tailored_summary)
        
        # Add education (moved before experience)
        if parsed_resume.education:
            self._add_education_section(doc, parsed_resume.education)
        
        # Add experience
        if optimized_experiences:
            self._add_experience_section(doc, optimized_experiences)
        
        # Add projects (if available)
        if parsed_resume.projects:
            self._add_projects_section(doc, parsed_resume.projects[:4])  # Top 4 projects
        
        # Add skills (organized by category)
        if selected_skills:
            self._add_skills_section(doc, selected_skills, parsed_resume.skills)
        
        # Generate filename
        filename = self._generate_filename(company_name, job_title)
        output_path = self.output_dir / filename
        
        # Save document
        doc.save(str(output_path))
        logger.info(f"DOCX saved: {output_path}")
        
        return str(output_path)
    
    def _add_header(self, doc: Document, contact_info):
        """Add resume header with name and contact info - Professional format"""
        # Name (large, bold, centered)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run((contact_info.name or "Professional").upper())
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_para.paragraph_format.space_after = Pt(4)
        
        # Contact info (centered, using ⋄ separator)
        contact_parts = []
        if contact_info.phone:
            contact_parts.append(contact_info.phone)
        if contact_info.location:
            contact_parts.append(contact_info.location)
        if contact_info.email:
            contact_parts.append(contact_info.email)
        if contact_info.linkedin:
            # Clean LinkedIn URL
            linkedin = contact_info.linkedin.replace('https://', '').replace('http://', '').replace('www.', '')
            contact_parts.append(linkedin)
        if contact_info.github:
            # Clean GitHub URL
            github = contact_info.github.replace('https://', '').replace('http://', '').replace('www.', '')
            contact_parts.append(github)
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" ⋄ ".join(contact_parts))
            contact_run.font.size = Pt(10)
            contact_para.paragraph_format.space_after = Pt(12)
    
    def _add_section(self, doc: Document, title: str, content: str):
        """Add a section with title and content - Professional format"""
        # Section title (uppercase, bold, no underline)
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title.upper())
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_para.paragraph_format.space_before = Pt(8)
        title_para.paragraph_format.space_after = Pt(6)
        
        # Content (justified)
        content_para = doc.add_paragraph(content)
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content_para.paragraph_format.space_after = Pt(10)
        content_para.paragraph_format.line_spacing = 1.15
    
    def _add_experience_section(
        self,
        doc: Document,
        experiences: List[Tuple[Experience, List[str]]]
    ):
        """Add experience section with optimized bullets - Professional format"""
        # Section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("EXPERIENCE")
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_para.paragraph_format.space_before = Pt(8)
        title_para.paragraph_format.space_after = Pt(6)
        
        # Add each experience
        for exp, bullets in experiences:
            # Create a table with 2 columns for title/dates alignment
            exp_table = doc.add_table(rows=1, cols=2)
            exp_table.autofit = False
            exp_table.allow_autofit = False
            
            # Left cell: Job title (bold)
            left_cell = exp_table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            left_run = left_para.add_run(exp.title)
            left_run.bold = True
            left_run.font.size = Pt(10.5)
            
            # Right cell: Dates (right-aligned)
            right_cell = exp_table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_str = f"{exp.start_date or ''} - {exp.end_date or 'Present'}"
            right_run = right_para.add_run(date_str)
            right_run.font.size = Pt(10.5)
            
            # Remove table borders
            for row in exp_table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(
                        doc._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
                    )
            
            # Company and location on next line
            company_para = doc.a - Professional format"""
        # Section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("EDUCATION")
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_para.paragraph_format.space_before = Pt(8)
        title_para.paragraph_format.space_after = Pt(6)
        
        # Add each degree
        for edu in education:
            # Create table for degree/date alignment
            edu_table = doc.add_table(rows=1, cols=2)
            edu_table.autofit = False
            edu_table.allow_autofit = False
            
            # Left cell: Degree and field
            left_cell = edu_table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            degree_text = edu.degree
            if edu.field:
                degree_text += f" in {edu.field}"
            degree_text += f", {edu.institution}"
            left_run = left_para.add_run(degree_text)
            left_run.bold = True
            left_run.font.size = Pt(10.5)
            
            # Right cell: Graduation date (right-aligned)
            right_cell = edu_table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if edu.graduation_date:
                right_run = right_para.add_run(edu.graduation_date)
                right_run.font.size = Pt(10.5)
            
            # Remove table borders
            for row in edu_table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(
                        doc._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
                    )
            
            # GPA on next line (always show if available)
            if edu.gpa:
                gpa_para = doc.add_paragraph()
                gpa_text = edu.gpa if 'GPA' in edu.gpa or 'CGPA' in edu.gpa else f"CGPA: {edu.gpa}"
                gpa_run = gpa_para.add_run(gpa_text)
                gpa_run.font.size = Pt(10)
                gpa_para.paragraph_format.space_before = Pt(0)
                gpa_para.paragraph_format.space_after = Pt(8)
            else:
                # Add spacing if no GPA
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(4
            # Institution and dates
            inst_parts = [edu.institution]
            if edu.graduation_date:
                inst_parts.append(edu.graduation_date)
            
            inst_para = doc.add_paragraph(" | ".join(inst_parts))
            inst_para.paragraph_format.space_before = Pt(0)
            inst_para.paragraph_format.space_after = Pt(6)
            
            # GPA if noteworthy
            if edu.gpa:
                try:
                    # Extract numeric value from GPA string (handles "3.8", "CGPA: 3.8", "3.8/4.0", etc.)
                    gpa_text = edu.gpa.replace('CGPA:', '').replace('GPA:', '').strip()
                    gpa_numeric = float(gpa_text.split()[0].split('/')[0])
                    
                    if gpa_numeric >= 3.5:
                        gpa_para = doc.add_paragraph(f"GPA: {edu.gpa}")
                        gpa_para.paragraph_format.left_indent = Inches(0.25)
                except (ValueError, IndexError):
                    # If GPA can't be parsed, skip it
                    logger.debug(f"Could not parse GPA: {edu.gpa}")
    
    def _add_projects_section(self, doc: Document, projects: List[Dict[str, str]]):
        """Add projects section - Professional format"""
        # Section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("PROJECTS")
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_para.paragraph_format.space_before = Pt(8)
        title_para.paragraph_format.space_after = Pt(6)
        
        # Add each project
        for project in projects:
            proj_para = doc.add_paragraph()
            
            # Project name (bold)
            name = project.get('name', 'Project')
            name_run = proj_para.add_run(name + '. ')
            name_run.bold = True
            name_run.font.size = Pt(10.5)
            
            # Project description (same line, not bold)
            description = project.get('description', '')
            desc_run = proj_para.add_run(description)
            desc_run.font.size = Pt(10.5)
            
            proj_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            proj_para.paragraph_format.space_after = Pt(6)
            proj_para.paragraph_format.line_spacing = 1.15
    
    def _add_skills_section(self, doc: Document, selected_skills: List[str], all_skills: List[str]):
        """Add skills section organized by category - Professional format"""
        # Section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("SKILLS")
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_para.paragraph_format.space_before = Pt(8)
        title_para.paragraph_format.space_after = Pt(6)
        
        # Categorize skills
        categories = {
            'Programming Languages': [],
            'Frameworks / Tools': [],
            'Machine Learning & AI': [],
            'Database Management': [],
            'DevOps': []
        }
        
        # Keywords for categorization
        prog_langs = ['python', 'c++', 'java', 'javascript', 'c#', 'kotlin', 'c ', 'swift']
        frameworks = ['react', 'angular', 'node', 'flask', 'django', 'spring', 'tailwind', 'qt']
        ml_ai = ['tensorflow', 'pytorch', 'keras', 'scikit', 'llm', 'gpt', 'whisper', 'langchain', 'numpy', 'pandas']
        databases = ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'firebase']
        devops = ['docker', 'kubernetes', 'git', 'aws', 'azure', 'gcp']
        
        # Categorize each skill
        for skill in selected_skills[:25]:
            skill_lower = skill.lower()
            categorized = False
            
            if any(lang in skill_lower for lang in prog_langs):
                categories['Programming Languages'].append(skill)
            elif any(fw in skill_lower for fw in frameworks):
                categories['Frameworks / Tools'].append(skill)
            elif any(ml in skill_lower for ml in ml_ai):
                categories['Machine Learning & AI'].append(skill)
            elif any(db in skill_lower for db in databases):
                categories['Database Management'].append(skill)
            elif any(dv in skill_lower for dv in devops):
                categories['DevOps'].append(skill)
        
        # Add each category
        for category, skills in categories.items():
            if skills:
                cat_para = doc.add_paragraph()
                cat_run = cat_para.add_run(category + '\n')
                cat_run.bold = True
                cat_run.font.size = Pt(10.5)
                skills_run = cat_para.add_run(', '.join(skills))
                skills_run.font.size = Pt(10)
                cat_para.paragraph_format.space_after = Pt(4)
    
    def _generate_filename(self, company_name: str, job_title: str) -> str:
        """
        Generate filename: Resume_[Company]_[Role].docx
        
        Args:
            company_name: Company name
            job_title: Job title
        
        Returns:
            Sanitized filename
        """
        # Sanitize strings (remove special characters)
        company_clean = re.sub(r'[^\w\s-]', '', company_name)
        company_clean = re.sub(r'[-\s]+', '_', company_clean)
        
        job_clean = re.sub(r'[^\w\s-]', '', job_title)
        job_clean = re.sub(r'[-\s]+', '_', job_clean)
        
        # Build filename
        filename = f"Resume_{company_clean}_{job_clean}.docx"
        
        # Ensure not too long (max 100 chars)
        if len(filename) > 100:
            filename = filename[:96] + ".docx"
        
        return filename
    
    def generate_multiple_versions(
        self,
        parsed_resume: ParsedResume,
        jobs: List[Dict[str, str]],
        batch_size: int = 5
    ) -> List[str]:
        """
        Generate multiple tailored resumes for different jobs
        
        Args:
            parsed_resume: Parsed master resume
            jobs: List of job dicts with 'description', 'company', 'title'
            batch_size: Number of resumes to generate
        
        Returns:
            List of generated file paths
        """
        logger.info(f"Generating {len(jobs)} tailored resumes")
        
        generated_files = []
        for job in jobs[:batch_size]:
            try:
                file_path = self.generate_tailored_resume(
                    parsed_resume=parsed_resume,
                    job_description=job['description'],
                    company_name=job['company'],
                    job_title=job['title']
                )
                generated_files.append(file_path)
            except Exception as e:
                logger.error(f"Failed to generate resume for {job['company']}: {e}")
        
        logger.info(f"Successfully generated {len(generated_files)} resumes")
        return generated_files


if __name__ == "__main__":
    # Example usage
    from parsers.resume_parser import ResumeParser
    
    # Parse master resume
    parser = ResumeParser()
    # parsed_resume = parser.parse_resume("path/to/resume.pdf")
    
    # Generate tailored resume
    generator = ResumeGenerator()
    # output = generator.generate_tailored_resume(
    #     parsed_resume=parsed_resume,
    #     job_description="Senior Python Developer...",
    #     company_name="Tech Corp",
    #     job_title="Senior Python Developer"
    # )
    
    print("Resume generator ready!")
