"""
Resume Parser Module
Supports PDF and DOCX formats with LLM-based section extraction
"""

import os
import re
from typing import Dict, List, Optional, Union
from pathlib import Path
import json

# Document processing
import pdfplumber
from docx import Document as DocxDocument

# LLM for intelligent parsing
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from pydantic import BaseModel, Field

# Logging
from utils.logger import logger

logger 


# Pydantic models for structured output
class ContactInfo(BaseModel):
    """Contact information from resume"""
    name: Optional[str] = Field(None, description="Full name")
    email: Optional[str] = Field(None, description="Email address")
    phone: Optional[str] = Field(None, description="Phone number")
    location: Optional[str] = Field(None, description="City, State or location")
    linkedin: Optional[str] = Field(None, description="LinkedIn profile URL")
    github: Optional[str] = Field(None, description="GitHub profile URL")
    website: Optional[str] = Field(None, description="Personal website")


class Experience(BaseModel):
    """Work experience entry"""
    title: str = Field(..., description="Job title")
    company: str = Field(..., description="Company name")
    location: Optional[str] = Field(None, description="Job location")
    start_date: Optional[str] = Field(None, description="Start date")
    end_date: Optional[str] = Field(None, description="End date (or 'Present')")
    description: List[str] = Field(default_factory=list, description="Bullet points describing responsibilities and achievements")


class Education(BaseModel):
    """Education entry"""
    degree: str = Field(..., description="Degree name (e.g., 'Bachelor of Science')")
    field: Optional[str] = Field(None, description="Field of study")
    institution: str = Field(..., description="School/University name")
    location: Optional[str] = Field(None, description="School location")
    graduation_date: Optional[str] = Field(None, description="Graduation date")
    gpa: Optional[str] = Field(None, description="GPA if mentioned")


class ParsedResume(BaseModel):
    """Complete parsed resume structure"""
    contact_info: ContactInfo = Field(..., description="Contact information")
    summary: Optional[str] = Field(None, description="Professional summary or objective")
    experience: List[Experience] = Field(default_factory=list, description="Work experience entries")
    education: List[Education] = Field(default_factory=list, description="Education entries")
    skills: List[str] = Field(default_factory=list, description="List of skills")
    certifications: List[str] = Field(default_factory=list, description="Certifications")
    projects: List[Dict[str, str]] = Field(default_factory=list, description="Projects with name and description")
    awards: List[str] = Field(default_factory=list, description="Awards and honors")
    raw_text: str = Field(..., description="Raw extracted text from resume")


class ResumeParser:
    """
    Parser for extracting structured data from resumes (PDF and DOCX)
    Uses LLM for intelligent section identification and extraction
    """
    
    def __init__(self, model: str = "gpt-4o-mini", temperature: float = 0.0):
        """
        Initialize resume parser
        
        Args:
            model: OpenAI model to use for parsing
            temperature: Model temperature (0.0 for deterministic)
        """
        self.llm = ChatOpenAI(
            model=model,
            temperature=temperature,
            openai_api_key=os.getenv("OPENAI_API_KEY")
        )
        logger.info(f"ResumeParser initialized with model: {model}")
    
    def parse_file(self, file_path: Union[str, Path]) -> ParsedResume:
        """
        Parse resume file (PDF or DOCX)
        
        Args:
            file_path: Path to resume file
            
        Returns:
            ParsedResume object with structured data
            
        Raises:
            ValueError: If file format not supported
            FileNotFoundError: If file doesn't exist
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        logger.info(f"Parsing resume: {file_path.name}")
        
        # Extract text based on file type
        if file_path.suffix.lower() == '.pdf':
            raw_text = self._extract_pdf_text(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            raw_text = self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        if not raw_text.strip():
            raise ValueError("No text could be extracted from the resume")
        
        logger.info(f"Extracted {len(raw_text)} characters from resume")
        
        # Use LLM to parse structured data
        parsed_resume = self._parse_with_llm(raw_text)
        
        logger.info("Resume parsing complete")
        return parsed_resume
    
    def _extract_pdf_text(self, pdf_path: Path) -> str:
        """
        Extract text from PDF file
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text_parts = []
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    logger.debug(f"Extracted text from page {page_num}")
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
    
    def _extract_docx_text(self, docx_path: Path) -> str:
        """
        Extract text from DOCX file
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        text_parts = []
        
        try:
            doc = DocxDocument(docx_path)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def _parse_with_llm(self, raw_text: str) -> ParsedResume:
        """
        Use LLM to parse resume into structured format
        
        Args:
            raw_text: Raw resume text
            
        Returns:
            ParsedResume object
        """
        # Create JSON output parser with the schema
        parser = JsonOutputParser(pydantic_object=ParsedResume)
        
        # Create prompt for LLM
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert resume parser. Extract structured information from the provided resume text.

Instructions:
1. Extract contact information (name, email, phone, location, LinkedIn, GitHub, website)
2. Identify and extract professional summary/objective
3. Parse work experience with job titles, companies, dates, and bullet points
4. Extract education details with degrees, schools, and dates
5. List all technical and soft skills mentioned
6. Extract certifications, projects, and awards if present
7. Be thorough but only include information actually present in the resume
8. For dates, preserve the original format from the resume

{format_instructions}"""),
            ("user", "Resume text:\n\n{resume_text}")
        ])
        
        # Create chain
        chain = prompt | self.llm | parser
        
        try:
            # Parse resume
            result = chain.invoke({
                "resume_text": raw_text,
                "format_instructions": parser.get_format_instructions()
            })
            
            # Add raw text to result
            result['raw_text'] = raw_text
            
            # Create ParsedResume object
            parsed_resume = ParsedResume(**result)
            
            logger.info(f"Parsed resume - Contact: {parsed_resume.contact_info.name}, "
                       f"Experience: {len(parsed_resume.experience)} entries, "
                       f"Education: {len(parsed_resume.education)} entries, "
                       f"Skills: {len(parsed_resume.skills)}")
            
            return parsed_resume
        
        except Exception as e:
            logger.error(f"Error parsing resume with LLM: {e}")
            # Return minimal ParsedResume with just raw text
            return ParsedResume(
                contact_info=ContactInfo(),
                raw_text=raw_text
            )
    
    def save_parsed_resume(self, parsed_resume: ParsedResume, output_path: Union[str, Path]):
        """
        Save parsed resume to JSON file
        
        Args:
            parsed_resume: ParsedResume object
            output_path: Path to save JSON file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(parsed_resume.model_dump(), f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved parsed resume to: {output_path}")
    
    def extract_contact_info_simple(self, text: str) -> Dict[str, str]:
        """
        Simple regex-based contact extraction (fallback method)
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary with contact information
        """
        contact = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact['email'] = email_match.group(0)
        
        # Phone (various formats)
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact['phone'] = phone_match.group(0)
        
        # LinkedIn
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = f"https://{linkedin_match.group(0)}"
        
        # GitHub
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact['github'] = f"https://{github_match.group(0)}"
        
        return contact


# Convenience function
def parse_resume(file_path: Union[str, Path], model: str = "gpt-4o-mini") -> ParsedResume:
    """
    Convenience function to parse a resume file
    
    Args:
        file_path: Path to resume file (PDF or DOCX)
        model: OpenAI model to use
        
    Returns:
        ParsedResume object
    """
    parser = ResumeParser(model=model)
    return parser.parse_file(file_path)


if __name__ == "__main__":
    # Example usage
    from dotenv import load_dotenv
    load_dotenv()
    
    # Test with a sample resume (you'll need to create this)
    sample_resume_path = Path("data/sample_resumes/john_doe_resume.pdf")
    
    if sample_resume_path.exists():
        print("📄 Parsing Resume...")
        parsed = parse_resume(sample_resume_path)
        
        print(f"\n✅ Resume Parsed Successfully!\n")
        print(f"Name: {parsed.contact_info.name}")
        print(f"Email: {parsed.contact_info.email}")
        print(f"Phone: {parsed.contact_info.phone}")
        print(f"\nExperience Entries: {len(parsed.experience)}")
        print(f"Education Entries: {len(parsed.education)}")
        print(f"Skills: {len(parsed.skills)}")
        
        # Save to JSON
        output_path = Path("output/parsed_resumes/john_doe_parsed.json")
        parser = ResumeParser()
        parser.save_parsed_resume(parsed, output_path)
        print(f"\n💾 Saved to: {output_path}")
    else:
        print(f"⚠️  Sample resume not found: {sample_resume_path}")
        print("Create a sample resume to test the parser")
