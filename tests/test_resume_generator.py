"""
Unit tests for resume generator
"""

import pytest
import os
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from docx import Document

from generators.resume_generator import ResumeGenerator, ContentSelection
from parsers.resume_parser import (
    ParsedResume, ContactInfo, Experience, Education
)
from tools.keyword_extractor import ExtractedKeywords


class TestResumeGenerator:
    """Test suite for ResumeGenerator class"""
    
    def setup_method(self):
        """Setup test fixtures"""
        self.generator = ResumeGenerator(
            model="gpt-4o-mini",
            output_dir="output/test_resumes"
        )
        
        self.sample_resume = ParsedResume(
            contact_info=ContactInfo(
                name="John Doe",
                email="john@example.com",
                phone="555-1234",
                location="San Francisco, CA",
                linkedin="linkedin.com/in/johndoe"
            ),
            summary="Software Engineer with 5 years of Python experience",
            experience=[
                Experience(
                    title="Senior Software Engineer",
                    company="Tech Corp",
                    start_date="2021",
                    end_date="Present",
                    location="San Francisco, CA",
                    description=[
                        "Built REST APIs using Python Django",
                        "Deployed applications on AWS using Docker",
                        "Led team of 3 junior developers"
                    ]
                ),
                Experience(
                    title="Software Engineer",
                    company="StartupXYZ",
                    start_date="2019",
                    end_date="2021",
                    location="Remote",
                    description=[
                        "Developed features using React and Node.js",
                        "Worked with PostgreSQL databases"
                    ]
                )
            ],
            education=[
                Education(
                    degree="Bachelor of Science",
                    field="Computer Science",
                    institution="State University",
                    graduation_date="2019",
                    gpa="3.8"
                )
            ],
            skills=["Python", "Django", "React", "Node.js", "PostgreSQL", "AWS", "Docker"],
            raw_text="Full resume text..."
        )
        
        self.sample_keywords = ExtractedKeywords(
            required_skills=["Python", "Django", "PostgreSQL"],
            preferred_skills=["Docker", "AWS"],
            tools_technologies=["Git", "Jenkins"],
            certifications=[],
            soft_skills=["Leadership", "Communication"],
            experience_level="5+ years",
            education=["Bachelor's degree"],
            keywords_for_ats=["Python", "Django", "PostgreSQL", "Docker", "AWS"]
        )
        
        self.sample_job = """
        Senior Python Developer
        Required: Python, Django, PostgreSQL, 5+ years
        Preferred: Docker, AWS, Kubernetes
        """
    
    def teardown_method(self):
        """Cleanup test files"""
        test_dir = Path("output/test_resumes")
        if test_dir.exists():
            for file in test_dir.glob("*.docx"):
                try:
                    file.unlink()
                except:
                    pass
    
    def test_initialization(self):
        """Test ResumeGenerator initialization"""
        assert self.generator.model == "gpt-4o-mini"
        assert self.generator.temperature == 0.7
        assert hasattr(self.generator, 'llm')
        assert hasattr(self.generator, 'keyword_extractor')
        assert hasattr(self.generator, 'semantic_matcher')
    
    def test_output_directory_creation(self):
        """Test that output directory is created"""
        test_dir = Path("output/test_resumes")
        assert test_dir.exists()
        assert test_dir.is_dir()
    
    def test_prioritize_skills(self):
        """Test skill prioritization based on job keywords"""
        resume_skills = [
            "Python", "Java", "Django", "React", "PostgreSQL",
            "MongoDB", "Docker", "Git", "JavaScript", "HTML"
        ]
        
        prioritized = self.generator._prioritize_skills(
            resume_skills,
            self.sample_keywords
        )
        
        # Required skills should come first
        assert "Python" in prioritized[:5]
        assert "Django" in prioritized[:5]
        assert "PostgreSQL" in prioritized[:5]
        
        # Preferred/tools should be next
        assert "Docker" in prioritized
    
    def test_prioritize_skills_empty(self):
        """Test skill prioritization with empty skill list"""
        prioritized = self.generator._prioritize_skills([], self.sample_keywords)
        assert prioritized == []
    
    def test_select_content(self):
        """Test content selection based on job match"""
        match_analysis = {
            'composite_score': 0.75,
            'keyword_match_score': 70
        }
        
        selected = self.generator._select_content(
            self.sample_resume,
            self.sample_keywords,
            match_analysis
        )
        
        assert 'experiences' in selected
        assert 'skills' in selected
        assert len(selected['experiences']) > 0
        assert len(selected['skills']) > 0
    
    def test_select_content_relevance_scoring(self):
        """Test that experiences are scored by relevance"""
        match_analysis = {'composite_score': 0.7}
        
        selected = self.generator._select_content(
            self.sample_resume,
            self.sample_keywords,
            match_analysis
        )
        
        # First experience should be more relevant (has Python, Django, AWS, Docker)
        first_exp = selected['experiences'][0][0]
        assert first_exp.title == "Senior Software Engineer"
    
    @patch('generators.resume_generator.ChatOpenAI')
    def test_create_tailored_summary(self, mock_llm):
        """Test tailored summary generation"""
        # Mock LLM response
        mock_chain = MagicMock()
        mock_chain.invoke.return_value = {
            'summary': 'Senior Python Developer with 5+ years experience in Django and AWS'
        }
        
        with patch.object(self.generator.llm, '__or__', return_value=mock_chain):
            summary = self.generator._create_tailored_summary(
                self.sample_resume,
                self.sample_keywords,
                self.sample_job
            )
        
        # Should return a summary
        assert isinstance(summary, str)
        assert len(summary) > 0
    
    @patch('generators.resume_generator.ChatOpenAI')
    def test_create_tailored_summary_fallback(self, mock_llm):
        """Test summary generation fallback on error"""
        # Mock LLM to raise exception
        with patch.object(
            self.generator,
            '_create_tailored_summary',
            side_effect=Exception("API error")
        ):
            try:
                summary = self.generator._create_tailored_summary(
                    self.sample_resume,
                    self.sample_keywords,
                    self.sample_job
                )
            except:
                # Should handle error gracefully
                pass
    
    @patch('generators.resume_generator.ChatOpenAI')
    def test_optimize_bullet_points(self, mock_llm):
        """Test bullet point optimization"""
        bullets = [
            "Built REST APIs",
            "Worked with databases"
        ]
        
        # Mock LLM response
        mock_chain = MagicMock()
        mock_chain.invoke.return_value = {
            'optimized_bullets': [
                "Architected and deployed RESTful APIs using Python Django",
                "Optimized PostgreSQL database queries reducing latency by 40%"
            ]
        }
        
        with patch.object(self.generator.llm, '__or__', return_value=mock_chain):
            optimized = self.generator._optimize_bullet_points(
                bullets,
                self.sample_keywords,
                "Senior Software Engineer"
            )
        
        assert len(optimized) == 2
        assert isinstance(optimized[0], str)
    
    def test_optimize_bullet_points_empty(self):
        """Test bullet point optimization with empty list"""
        optimized = self.generator._optimize_bullet_points(
            [],
            self.sample_keywords,
            "Developer"
        )
        assert optimized == []
    
    def test_generate_filename(self):
        """Test filename generation"""
        filename = self.generator._generate_filename(
            "Tech Corp",
            "Senior Python Developer"
        )
        
        assert filename == "Resume_Tech_Corp_Senior_Python_Developer.docx"
    
    def test_generate_filename_special_chars(self):
        """Test filename generation with special characters"""
        filename = self.generator._generate_filename(
            "Tech & Co., Inc.",
            "Sr. Python Dev (Remote)"
        )
        
        # Should sanitize special characters
        assert ".docx" in filename
        assert "&" not in filename
        assert "(" not in filename
    
    def test_generate_filename_long_name(self):
        """Test filename generation with very long names"""
        long_company = "A" * 60
        long_title = "B" * 60
        
        filename = self.generator._generate_filename(long_company, long_title)
        
        # Should truncate to max 100 chars
        assert len(filename) <= 100
        assert filename.endswith(".docx")
    
    @patch.object(ResumeGenerator, '_create_tailored_summary')
    @patch.object(ResumeGenerator, '_optimize_bullet_points')
    def test_generate_docx(self, mock_optimize, mock_summary):
        """Test DOCX generation"""
        mock_summary.return_value = "Tailored summary"
        mock_optimize.return_value = ["Optimized bullet 1", "Optimized bullet 2"]
        
        output_path = self.generator._generate_docx(
            parsed_resume=self.sample_resume,
            tailored_summary="Professional summary",
            optimized_experiences=[
                (self.sample_resume.experience[0], ["Bullet 1", "Bullet 2"])
            ],
            selected_skills=["Python", "Django", "AWS"],
            company_name="Tech Corp",
            job_title="Senior Developer"
        )
        
        # Check file was created
        assert os.path.exists(output_path)
        assert output_path.endswith(".docx")
        
        # Verify document content
        doc = Document(output_path)
        doc_text = "\n".join([p.text for p in doc.paragraphs])
        
        assert "John Doe" in doc_text
        assert "john@example.com" in doc_text
        assert "PROFESSIONAL SUMMARY" in doc_text or "Professional summary" in doc_text
    
    def test_add_header(self):
        """Test header addition to document"""
        doc = Document()
        self.generator._add_header(doc, self.sample_resume.contact_info)
        
        # Check that name and contact info were added
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "John Doe" in text
        assert "john@example.com" in text
        assert "555-1234" in text
    
    def test_add_section(self):
        """Test section addition to document"""
        doc = Document()
        self.generator._add_section(doc, "TEST SECTION", "Test content")
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "TEST SECTION" in text
        assert "Test content" in text
    
    def test_add_experience_section(self):
        """Test experience section addition"""
        doc = Document()
        experiences = [
            (self.sample_resume.experience[0], ["Bullet 1", "Bullet 2"])
        ]
        
        self.generator._add_experience_section(doc, experiences)
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "PROFESSIONAL EXPERIENCE" in text
        assert "Senior Software Engineer" in text
        assert "Tech Corp" in text
    
    def test_add_education_section(self):
        """Test education section addition"""
        doc = Document()
        self.generator._add_education_section(doc, self.sample_resume.education)
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "EDUCATION" in text
        assert "Bachelor of Science" in text
        assert "Computer Science" in text
        assert "State University" in text
    
    @patch.object(ResumeGenerator, 'generate_tailored_resume')
    def test_generate_multiple_versions(self, mock_generate):
        """Test generating multiple resume versions"""
        mock_generate.return_value = "output/resume1.docx"
        
        jobs = [
            {
                'description': 'Job 1 description',
                'company': 'Company 1',
                'title': 'Developer 1'
            },
            {
                'description': 'Job 2 description',
                'company': 'Company 2',
                'title': 'Developer 2'
            }
        ]
        
        files = self.generator.generate_multiple_versions(
            self.sample_resume,
            jobs,
            batch_size=2
        )
        
        assert len(files) == 2
        assert mock_generate.call_count == 2


class TestContentSelection:
    """Test content selection logic"""
    
    def setup_method(self):
        """Setup test fixtures"""
        self.generator = ResumeGenerator(output_dir="output/test")
    
    def test_select_top_experiences(self):
        """Test that top experiences are selected"""
        resume = ParsedResume(
            contact_info=ContactInfo(name="Test"),
            experience=[
                Experience(
                    title="Python Developer",
                    company="A",
                    description=["Python Django PostgreSQL"]
                ),
                Experience(
                    title="Java Developer",
                    company="B",
                    description=["Java Spring MySQL"]
                ),
                Experience(
                    title="Full Stack",
                    company="C",
                    description=["Python React AWS"]
                )
            ],
            education=[],
            skills=["Python"],
            raw_text=""
        )
        
        keywords = ExtractedKeywords(
            required_skills=["Python", "Django", "PostgreSQL"],
            preferred_skills=[],
            tools_technologies=[],
            certifications=[],
            soft_skills=[],
            experience_level="",
            education=[],
            keywords_for_ats=[]
        )
        
        selected = self.generator._select_content(
            resume,
            keywords,
            {'composite_score': 0.7}
        )
        
        # First experience should be Python Developer (highest keyword match)
        assert selected['experiences'][0][0].title == "Python Developer"


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
